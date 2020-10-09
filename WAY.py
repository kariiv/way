import pandas as pd
from collections import Counter
from random import shuffle, choice

file_name = "WAY"

# Need to edit WAY.csv file path
way = pd.read_csv("~/Desktop/" + file_name + ".csv", delimiter=',', encoding='utf-8', keep_default_na=False)

people = {}

limit = []
blacklist = []
errors = []
PASS_COUNT = 0
total_comments = 0
total_comments_on_docx = 0

for x, column in enumerate(way):
    if x == 0:  #Timestamp
        continue
    if "Unnamed:" in column:
        continue
    people[column] = []
    for y, row in enumerate(way[column]):
        if row and row not in ["-", "--"]:
            try:
                people[column].append(row)
                total_comments += 1
            except UnicodeEncodeError as ex:
                errors.append(row)

from docx import Document
from docx.shared import Pt

# Modify the name/text size

font_size_name = Pt(16)
font_size_text = Pt(12)

# Modify the name/text font family

font_name_name = "Calibri"
font_name_text = "Calibri"


blacklist_items = [""]
blacklist_items = [name.lower() for name in blacklist_items]

set_limit = 3

# nice_text_to_all = ["Rahulikku jõuluaega!", "", ""]

document = Document()
for k in sorted(people.keys(), key=lambda x: sum([len(comment) for comment in people[x]])): # Sort by text length
    if len(people[k]) < set_limit:
        limit.append(k + f"({len(people[k])})")
        continue
    if ":" in k:  # Removes the additional description ex:"President: " part from names
        name = k.split(":")[1].strip()
    else:
        name = k
    
    if name.lower() in blacklist_items:
        blacklist.append(name)
        continue
    
    PASS_COUNT += 1
    
    run = document.add_heading(name, level=0).add_run()
    run.font.name = font_name_name
    run.font.size = font_size_name
    run.font.bold = True

    # add_text = choice(nice_text_to_all)
    # if add_text:
    #     people[k].append(add_text)

    [shuffle(people[k]) for i in range(20)]  # Shuffle the results to make it more anonymous
    for line in people[k]:
        total_comments_on_docx += 1
        run = document.add_paragraph(line)
        # run.font.name = font_name_text
        # run.font.size = font_size_text
    document.add_page_break()

# Change the saving path 
document.save('/home/syrstyle/Desktop/' + file_name + ".docx")

counts = Counter(sorted([len(v) for k, v in people.items()]))
print(counts)

removed_lists = [limit, blacklist]

print("Totel People: ", len(people))
print(f"People In docx({PASS_COUNT})")
print(f"Total comments({total_comments})")
print(f"Comments in docx({total_comments_on_docx})")
print(f"ERRORS({len(errors)}):  {errors}")
print(f"LIMIT({len(limit)})", limit)
print(f"BLACKLIST({len(blacklist)}):", blacklist)
print(f"REMOVED({sum([len(_list) for _list in removed_lists])})")
